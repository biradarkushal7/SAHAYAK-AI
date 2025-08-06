import os
import re
import io # New import for in-memory file handling
 
from docx import Document
from docx.shared import Pt
from google.adk.agents import Agent # Keep this import if it's used elsewhere in your original file
 
# NEW IMPORTS FOR GCS
from google.cloud import storage
import datetime # For unique filenames
import mimetypes # For content type detection
 
# --- NEW: Define your GCS bucket for generated files ---
# !! IMPORTANT: Replace "your-generated-files-bucket-name" with your actual GCS bucket name !!
GENERATED_GCS_BUCKET_NAME = "gs://adk-test-bucket-005"
 
school_guardrails = """
## School Environment Guardrails
- Do NOT use or allow any profanity, hate speech, or inappropriate language.
- Do NOT generate or reference any content that is violent, sexual, discriminatory, or otherwise inappropriate for a school setting.
- Always maintain a respectful, positive, and encouraging tone suitable for children and students.
- Do NOT provide or suggest any unsafe, illegal, or harmful activities.
- If a user asks for something inappropriate, politely refuse and remind them this is a school environment.
"""
local_language_policy = """
## Local Language Policy
- If the user asks a question in a local (non-English) language, reply in the same language but type your response in English script (transliteration).
- Always match the user's language and tone, but keep the script in English.
"""
 
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
 
def set_run(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
 
def add_formatted_text(paragraph, text):
    pattern = re.compile(r'(\^\{.*?\}|\^[a-zA-Z0-9]|_\{.*?\}|_[a-zA-Z0-9])')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith('^'):
            val = token[2:-1] if token.startswith('^{') else token[1]
            run = paragraph.add_run(val)
            run.font.superscript = True
        elif token.startswith('_'):
            val = token[2:-1] if token.startswith('_{') else token[1]
            run = paragraph.add_run(val)
            run.font.subscript = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
 
# MODIFIED: Added user_id parameter
def save_answers_to_word(worksheet_answers: str, user_id: str) -> dict:
    """
    This tool converts the answers of the worksheet to a word file, uploads it to GCS,
    and returns a downloadable link.
 
    Args:
        worksheet_answers (str): contents of the worksheet answers
        user_id (str): The ID of the user, used for GCS folder structure.
 
    Returns:
        dict: status and downloadable_file_url
    """
    if not worksheet_answers:
        return {"status": "Missing worksheet_answers in context"}
    if not user_id: # Ensure user_id is provided
        return {"status": "Missing user_id for GCS upload"}
 
    lines = [line.strip() for line in worksheet_answers.strip().splitlines() if line.strip()]
    doc = Document()
 
    p = doc.add_paragraph()
    r = p.add_run("Worksheet Answers")
    set_run(r, size=14, bold=True)
 
    for line in lines:
        match_q = re.match(r'^(\d+)\.\s*(.+)$', line)
        if match_q:
            para = doc.add_paragraph()
            add_formatted_text(para, f"{match_q.group(1)}. {match_q.group(2)}")
        else:
            para = doc.add_paragraph()
            add_formatted_text(para, line)
 
    # --- NEW: Save document to an in-memory buffer ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0) # Rewind the stream to the beginning
 
    # --- NEW: Upload to GCS ---
    storage_client = storage.Client()
    bucket = storage_client.bucket(GENERATED_GCS_BUCKET_NAME)
 
    # Generate a unique filename for GCS
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    # Using a hash of answers for uniqueness is fine, but combine with timestamp for better uniqueness
    unique_id = hash(worksheet_answers) % 100000
    gcs_object_name = f"{user_id}/generated/answers_{timestamp}_{unique_id}.docx"
 
    blob = bucket.blob(gcs_object_name)
 
    # Upload the in-memory file stream
    blob.upload_from_file(file_stream, content_type=DOCX_MIME)
 
    # Make the blob publicly accessible if desired (for direct download links)
    # WARNING: This makes the file publicly readable. Adjust permissions as needed.
    # For signed URLs (more secure), see notes below.
    blob.make_public()
 
    # Construct the downloadable URL
    downloadable_url = blob.public_url
 
    return {"status": "success", "downloadable_file_url": downloadable_url}
 
 
root_agent = Agent(
    name="answering_agent",
    model="gemini-2.0-flash",
    description="Solve worksheet questions and save them to a Word file in the worksheets/ directory.",
    instruction=f"""
        You are an expert in solving worksheet questions.
        Do not display the worksheet questions or answers in your response.
        For each worksheet answer, include the step-by-step solution as plain text directly after the answer number, without any special formatting.
        For example, for the question 'Find the roots of the quadratic equation x2 - 4x + 4 = 0.', the output should be:
        1) (x - 2)(x - 2) = 0\n   x - 2 = 0 | x - 2 = 0\n   x = 2 | x = 2\n
        Only call the `save_answers_to_word` tool for saving the solution into a word file and return a confirmation message like:
        'Your worksheet has been solved successfully. The answers are available for download here: <downloadable_file_url>.'
        When calling `save_answers_to_word`, ensure you pass the `user_id` from the current conversation context.
{school_guardrails}
{local_language_policy}
    """,
    tools=[save_answers_to_word],
)
 
 
 
# import os
# import re
 
# from docx import Document
# from docx.shared import Pt
# from google.adk.agents import Agent
# school_guardrails = """
# ## School Environment Guardrails
# - Do NOT use or allow any profanity, hate speech, or inappropriate language.
# - Do NOT generate or reference any content that is violent, sexual, discriminatory, or otherwise inappropriate for a school setting.
# - Always maintain a respectful, positive, and encouraging tone suitable for children and students.
# - Do NOT provide or suggest any unsafe, illegal, or harmful activities.
# - If a user asks for something inappropriate, politely refuse and remind them this is a school environment.
# """
# local_language_policy = """
# ## Local Language Policy
# - If the user asks a question in a local (non-English) language, reply in the same language but type your response in English script (transliteration).
# - Always match the user's language and tone, but keep the script in English.
# """
 
# DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
 
# def set_run(run, size=12, bold=False):
#     run.font.size = Pt(size)
#     run.bold = bold
#     run.font.name = "Calibri"
 
# def add_formatted_text(paragraph, text):
#     # Only handle superscript and subscript
#     pattern = re.compile(r'(\^\{.*?\}|\^[a-zA-Z0-9]|_\{.*?\}|_[a-zA-Z0-9])')
#     pos = 0
#     for m in pattern.finditer(text):
#         if m.start() > pos:
#             paragraph.add_run(text[pos:m.start()])
#         token = m.group()
#         if token.startswith('^'):
#             val = token[2:-1] if token.startswith('^{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.superscript = True
#         elif token.startswith('_'):
#             val = token[2:-1] if token.startswith('_{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.subscript = True
#         pos = m.end()
#     if pos < len(text):
#         paragraph.add_run(text[pos:])
 
# def save_answers_to_word(worksheet_answers: str) -> dict:
#     """
#     This tool converts the answers of the worksheet to a word file in worksheets/ directory
 
#     Args:
#         worksheet_answers (str): contents of the worksheet answers
 
#     Returns:
#         dict: status and worksheet_file path
#     """
#     if not worksheet_answers:
#         return {"status": "Missing worksheet_answers in context"}
 
#     lines = [line.strip() for line in worksheet_answers.strip().splitlines() if line.strip()]
#     doc = Document()
   
#     # Add title
#     p = doc.add_paragraph()
#     r = p.add_run("Worksheet Answers")
#     set_run(r, size=14, bold=True)
   
#     # Process each answer line
#     for line in lines:
#         # Check if it's a numbered answer (e.g., "1. ...")
#         match_q = re.match(r'^(\d+)\.\s*(.+)$', line)
#         if match_q:
#             para = doc.add_paragraph()
#             add_formatted_text(para, f"{match_q.group(1)}. {match_q.group(2)}")
#         else:
#             # If not numbered, add as regular paragraph
#             para = doc.add_paragraph()
#             add_formatted_text(para, line)
 
#     target_path = r'C:\Users\birad\Downloads\Hack2skill_working_db\multi_agent\manager\sub_agents\worksheet_solver\worksheets'
   
# #     # Create the directory
#     os.makedirs(target_path, exist_ok=True)
#     filename = f"{target_path}/answers_{hash(worksheet_answers) % 100000}.docx"
#     doc.save(filename)
#     return {"status": "success", "worksheet_file": filename}
 
 
# root_agent = Agent(
#     name="answering_agent",
#     model="gemini-2.0-flash",
#     description="Solve worksheet questions and save them to a Word file in the worksheets/ directory.",
#     instruction=f"""
#         You are an expert in solving worksheet questions.
#         Do not display the worksheet questions or answers in your response.
#         For each worksheet answer, include the step-by-step solution as plain text directly after the answer number, without any special formatting.
#         For example, for the question 'Find the roots of the quadratic equation x2 - 4x + 4 = 0.', the output should be:
#         1) (x - 2)(x - 2) = 0\n   x - 2 = 0 | x - 2 = 0\n   x = 2 | x = 2\n
#         Only call the `save_answers_to_word` tool for saving the solution into a word file and return a confirmation message like:
#         'Your worksheet has been solved successfully. The answers are available for download in the worksheets/ folder as <filename>.docx.'
# {school_guardrails}
# {local_language_policy}
#     """,
#     tools=[save_answers_to_word],
# )
 
 
 