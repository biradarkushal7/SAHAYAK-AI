import os
import re
import io # New import for in-memory file handling
 
from docx import Document
from docx.shared import Pt
from google.adk.agents import Agent # Keep this if it's used elsewhere
 
# NEW IMPORTS FOR GCS
from google.cloud import storage
import datetime # For unique filenames
import mimetypes # For content type detection
 
# --- NEW: Define your GCS bucket for generated files ---
# !! IMPORTANT: Replace "your-generated-files-bucket-name" with your actual GCS bucket name !!
GENERATED_GCS_BUCKET_NAME = "gs://adk-test-bucket-005"
 
school_guardrails = """
## School Environment Guardrails
- Do NOT use or allow any profanity, hate speech, or inappropriate language.
- Do NOT generate or reference any content that is violent, sexual, discriminatory, or otherwise inappropriate for a school setting.
- Always maintain a respectful, positive, and encouraging tone suitable for children and students.
- Do NOT provide or suggest any unsafe, illegal, or harmful activities.
- If a user asks for something inappropriate, politely refuse and remind them this is a school environment.
"""
local_language_policy = """
## Local Language Policy
- If the user asks a question in a local (non-English) language, reply in the same language but type your response in English script (transliteration).
- Always match the user's language and tone, but keep the script in English.
"""
def set_run(run, size=12, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = "Calibri"
 
def add_formatted_text(paragraph, text):
    import re
    pattern = re.compile(r'(\^\{.*?\}|\^[a-zA-Z0-9]|_\{.*?\}|_[a-zA-Z0-9])')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group()
        if token.startswith('^'):
            val = token[2:-1] if token.startswith('^{') else token[1]
            run = paragraph.add_run(val)
            run.font.superscript = True
        elif token.startswith('_'):
            val = token[2:-1] if token.startswith('_{') else token[1]
            run = paragraph.add_run(val)
            run.font.subscript = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
 
# MODIFIED: Added user_id parameter
def create_worksheet_tool(worksheet_text: str, user_id: str) -> dict:
    """
    This tool converts the contents of the worksheet to a Word file, uploads it to GCS,
    and returns a downloadable link.
 
    Args:
        worksheet_text (str): contents of the worksheet
        user_id (str): The ID of the user, used for GCS folder structure.
 
    Returns:
        dict: status and downloadable_file_url
    """
    if not worksheet_text:
        return {"status": "Missing worksheet_text in context"}
    if not user_id:
        return {"status": "Missing user_id for GCS upload"}
 
    lines = [line.strip() for line in worksheet_text.strip().splitlines() if line.strip()]
    doc = Document()
    section = None
    match_left, match_right = [], []
 
    i = 0
    while i < len(lines):
        line = lines[i]
        if line in ["Fill in the Blanks", "Match the Following", "Multiple Choice Questions", "Short Answer Type Questions", "True or False","Solve These Problems"]:
            section = line
            p = doc.add_paragraph()
            r = p.add_run(section)
            set_run(r, size=14, bold=True)
            i += 1
            continue
 
        if section == "Match the Following":
            m_left = re.match(r'^(\d+)\.\s*(.+)$', line)
            m_right = re.match(r'^([a-z])\)\s*(.+)$', line)
            if m_left:
                match_left.append(m_left.group(2))
            elif m_right:
                match_right.append(m_right.group(2))
            if (i + 1 == len(lines) or lines[i + 1] in ["Multiple Choice Questions", "Short Answer Type Questions", "True or False"]):
                table = doc.add_table(rows=max(len(match_left), len(match_right)), cols=2)
                table.style = "Table Grid"
                for idx in range(len(match_left)):
                    add_formatted_text(table.cell(idx, 0).paragraphs[0], f"{idx+1}. {match_left[idx]}")
                    if idx < len(match_right):
                        add_formatted_text(table.cell(idx, 1).paragraphs[0], f"{chr(97+idx)}) {match_right[idx]}")
                match_left, match_right = [], []
            i += 1
            continue
 
        if section == "Multiple Choice Questions":
            q_match = re.match(r'^(\d+)\.\s*(.+)$', line)
            if q_match:
                question = q_match.group(2)
                options = []
                j = i + 1
                while j < len(lines) and re.match(r'^[a-d]\)', lines[j]):
                    options.append(lines[j].strip())
                    j += 1
                para = doc.add_paragraph()
                add_formatted_text(para, f"{q_match.group(1)}. {question}")
                opt_line = "    ".join(options)
                if len(opt_line) < 70:
                    add_formatted_text(para, "\t" + opt_line)
                else:
                    for opt in options:
                        opt_p = doc.add_paragraph("\t")
                        add_formatted_text(opt_p, opt)
                i = j
                continue
 
        match_q = re.match(r'^(\d+)\.\s*(.+)$', line)
        if match_q:
            para = doc.add_paragraph()
            add_formatted_text(para, f"{match_q.group(1)}. {match_q.group(2)}")
        i += 1
 
    # --- REPLACED LOCAL SAVE WITH GCS UPLOAD ---
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
 
    storage_client = storage.Client()
    bucket = storage_client.bucket(GENERATED_GCS_BUCKET_NAME)
 
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = hash(worksheet_text) % 100000
    gcs_object_name = f"{user_id}/generated/worksheet_{timestamp}_{unique_id}.docx"
 
    blob = bucket.blob(gcs_object_name)
    blob.upload_from_file(file_stream, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    blob.make_public() # Makes the file publicly readable for a direct download link
 
    downloadable_url = blob.public_url
 
    return {"status": "success", "downloadable_file_url": downloadable_url}
 
 
root_agent = Agent(
    name="worksheet_content_agent",
    model="gemini-2.5-flash",
    description="Creates content to be populated in the worksheet.",
    instruction=f"""
You are an expert at creating content of worksheets for students provided the class, syllabus and concept for creating the worksheet.
 
The types of questions you must set will be as follows (these will apply for all subjects apart from Mathematics):
    - Fill in the Blanks
    - Match the Following
    - Multiple Choice Questions
    - Short Answer Type Questions
    - True or False
In case of Mathematics, do the following:
    - type of questions:"solve these problems" and the content should be in the format of a worksheet.
    - Create numerical problems of varying difficulties based on the syllabus and concept provided by the user
 
   
Formatting instructions for worksheet content:
    - For superscripts, use ^ (e.g., x^2; for multi-character superscripts, use curly brackets to enclose the entire superscript)
    - For subscripts, use _ (e.g., H_2; for multi-character subscripts, use curly brackets to enclose the entire subscript)
    - For equations, fractions, and roots, use plain text: e.g., 1/2, sqrt(x), x^2 + y^2 = 1
    - For math symbols, use Unicode if possible (e.g., ≤, ≥, π, ×, ÷, ∪, ∩, ∈, ∉, ∞, etc.)
    - For chemical symbols, use standard notation (e.g., H₂SO₄)
    - Do NOT use LaTeX, $...$, or \frac, \leq, etc.
    - Ensure formatting is consistent and clear for all questions
 
Keep the following things in mind while creating worksheets:
    - Make sure the class, syllabus and concept are mentioned before creating the worksheet. If any of these are missing, ask the user to mention them and do not proceed until they are properly passed.
    - If not mentioned by the user, you will ask them how many questions per question type they expect and if they don't have a preference, default to 5 questions per question type
    - You have to create questions only based on the specified class, syllabus and concept. No questions should appear beyond this constraint.
    - Try to include questions of all difficulties (Easy, Medium, Hard) without mentioning which level they are, within the above mentioned constraint.
    - The different question types will be grouped under their respective question type header names.
    For example, if the heading is 'Fill in the Blanks', the questions below it must be related to fill in the blanks.
 
Here's the instructions on how you should return the content of the worksheet:
    - Return only the question type headings followed by the respective questions under it. **DO NOT** send anything extra apart from it.
    - The output structure should be as demonstrated in the following example worksheet for physics:
        Example:
        Fill in the Blanks
        1. Electric current is measured in ____
        2. The SI unit of resistance is ____
 
        Match the Following
        3. Volt
        4. Ampere
        a) Unit of current
        b) Unit of potential difference
 
        Multiple Choice Questions
        6. Which instrument measures electric current?
        a) Voltmeter
        b) Ammeter
        c) Thermometer
        d) Barometer
 
        Short Answer Type Questions
        8. Define electric current
        9. What is Ohm's Law?
       
        True or False
        10. Ampere is the SI unit of resistance
        11. Current is directly proportional to Voltage if Resistance is constant
IMPORTANT:
    USE `create_worksheet_tool` to create the worksheet word file and return the dictionary (as JSON) containing the following keys:
    `response`: model response indicating the status of the worksheet creation
    `agent_called`: name of the agent that created the worksheet
    `filepath`: path to the created worksheet file
    When calling `create_worksheet_tool`, ensure you pass the `user_id` from the current conversation context.
    # **NEVER DISPLAY OR RETURN THE WORKSHEET FILE CONTENT TO THE USER**
{school_guardrails}
{local_language_policy}
 
"""
,
    tools=[create_worksheet_tool],
)
 
# import os
# import re
# from docx import Document
# from docx.shared import Pt
# from google.adk.agents import Agent
 
# school_guardrails = """
# ## School Environment Guardrails
# - Do NOT use or allow any profanity, hate speech, or inappropriate language.
# - Do NOT generate or reference any content that is violent, sexual, discriminatory, or otherwise inappropriate for a school setting.
# - Always maintain a respectful, positive, and encouraging tone suitable for children and students.
# - Do NOT provide or suggest any unsafe, illegal, or harmful activities.
# - If a user asks for something inappropriate, politely refuse and remind them this is a school environment.
# """
# local_language_policy = """
# ## Local Language Policy
# - If the user asks a question in a local (non-English) language, reply in the same language but type your response in English script (transliteration).
# - Always match the user's language and tone, but keep the script in English.
# """
# def set_run(run, size=12, bold=False):
#     run.font.size = Pt(size)
#     run.bold = bold
#     run.font.name = "Calibri"
 
# # Helper to format subscripts and superscripts in a run (no LaTeX, no $...$)
# def add_formatted_text(paragraph, text):
#     import re
#     # Only handle superscript and subscript
#     pattern = re.compile(r'(\^\{.*?\}|\^[a-zA-Z0-9]|_\{.*?\}|_[a-zA-Z0-9])')
#     pos = 0
#     for m in pattern.finditer(text):
#         if m.start() > pos:
#             paragraph.add_run(text[pos:m.start()])
#         token = m.group()
#         if token.startswith('^'):
#             val = token[2:-1] if token.startswith('^{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.superscript = True
#         elif token.startswith('_'):
#             val = token[2:-1] if token.startswith('_{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.subscript = True
#         pos = m.end()
#     if pos < len(text):
#         paragraph.add_run(text[pos:])
 
# def create_worksheet_tool(worksheet_text: str) -> dict:
#     """
#     This tool converts the contents of the worksheet to a word file
 
#     Args:
#         worksheet_text (str): contents of the worksheet
 
#     Returns:
#         dict:
#     """
#     if not worksheet_text:
#         return {"status": "Missing worksheet_text in context"}
 
#     lines = [line.strip() for line in worksheet_text.strip().splitlines() if line.strip()]
#     doc = Document()
#     section = None
#     match_left, match_right = [], []
 
#     i = 0
#     while i < len(lines):
#         line = lines[i]
#         # Section headers
#         if line in ["Fill in the Blanks", "Match the Following", "Multiple Choice Questions", "Short Answer Type Questions", "True or False","Solve These Problems"]:
#             section = line
#             p = doc.add_paragraph()
#             r = p.add_run(section)
#             set_run(r, size=14, bold=True)
#             i += 1
#             continue
 
#         if section == "Match the Following":
#             m_left = re.match(r'^(\d+)\.\s*(.+)$', line)
#             m_right = re.match(r'^([a-z])\)\s*(.+)$', line)
#             if m_left:
#                 match_left.append(m_left.group(2))
#             elif m_right:
#                 match_right.append(m_right.group(2))
#             if (i + 1 == len(lines) or lines[i + 1] in ["Multiple Choice Questions", "Short Answer Type Questions", "True or False"]):
#                 table = doc.add_table(rows=max(len(match_left), len(match_right)), cols=2)
#                 table.style = "Table Grid"
#                 for idx in range(len(match_left)):
#                     add_formatted_text(table.cell(idx, 0).paragraphs[0], f"{idx+1}. {match_left[idx]}")
#                     if idx < len(match_right):
#                         add_formatted_text(table.cell(idx, 1).paragraphs[0], f"{chr(97+idx)}) {match_right[idx]}")
#                 match_left, match_right = [], []
#             i += 1
#             continue
 
#         if section == "Multiple Choice Questions":
#             q_match = re.match(r'^(\d+)\.\s*(.+)$', line)
#             if q_match:
#                 question = q_match.group(2)
#                 options = []
#                 j = i + 1
#                 while j < len(lines) and re.match(r'^[a-d]\)', lines[j]):
#                     options.append(lines[j].strip())
#                     j += 1
#                 para = doc.add_paragraph()
#                 add_formatted_text(para, f"{q_match.group(1)}. {question}")
#                 opt_line = "    ".join(options)
#                 if len(opt_line) < 70:
#                     add_formatted_text(para, "\t" + opt_line)
#                 else:
#                     for opt in options:
#                         opt_p = doc.add_paragraph("\t")
#                         add_formatted_text(opt_p, opt)
#                 i = j
#                 continue
 
#         # Other questions
#         match_q = re.match(r'^(\d+)\.\s*(.+)$', line)
#         if match_q:
#             para = doc.add_paragraph()
#             add_formatted_text(para, f"{match_q.group(1)}. {match_q.group(2)}")
#         i += 1
 
#     base_path = r"C:\Users\birad\Downloads\Hack2skill_working_db\multi_agent\manager\sub_agents\worksheet"
#     target_path = os.path.join(base_path, "worksheets")
   
#     # Create the directory
#     os.makedirs(target_path, exist_ok=True)
#     filename = f"worksheets/worksheet_{hash(worksheet_text) % 100000}.docx"
#     doc.save(filename)
#     return {"status": "success", "worksheet_file": filename}
 
 
# root_agent = Agent(
#     name="worksheet_content_agent",
#     # https://ai.google.dev/gemini-api/docs/models
#     model="gemini-2.5-flash",
#     description="Creates content to be populated in the worksheet.",
#     instruction=f"""
# You are an expert at creating content of worksheets for students provided the class, syllabus and concept for creating the worksheet.
 
# The types of questions you must set will be as follows (these will apply for all subjects apart from Mathematics):
#     - Fill in the Blanks
#     - Match the Following
#     - Multiple Choice Questions
#     - Short Answer Type Questions
#     - True or False
# In case of Mathematics, do the following:
#     - type of questions:"solve these problems" and the content should be in the format of a worksheet.
#     - Create numerical problems of varying difficulties based on the syllabus and concept provided by the user
 
   
# Formatting instructions for worksheet content:
#     - For superscripts, use ^ (e.g., x^2; for multi-character superscripts, use curly brackets to enclose the entire superscript)
#     - For subscripts, use _ (e.g., H_2; for multi-character subscripts, use curly brackets to enclose the entire subscript)
#     - For equations, fractions, and roots, use plain text: e.g., 1/2, sqrt(x), x^2 + y^2 = 1
#     - For math symbols, use Unicode if possible (e.g., ≤, ≥, π, ×, ÷, ∪, ∩, ∈, ∉, ∞, etc.)
#     - For chemical symbols, use standard notation (e.g., H₂SO₄)
#     - Do NOT use LaTeX, $...$, or \frac, \leq, etc.
#     - Ensure formatting is consistent and clear for all questions
 
# Keep the following things in mind while creating worksheets:
#     - Make sure the class, syllabus and concept are mentioned before creating the worksheet. If any of these are missing, ask the user to mention them and do not proceed until they are properly passed.
#     - If not mentioned by the user, you will ask them how many questions per question type they expect and if they don't have a preference, default to 5 questions per question type
#     - You have to create questions only based on the specified class, syllabus and concept. No questions should appear beyond this constraint.
#     - Try to include questions of all difficulties (Easy, Medium, Hard) without mentioning which level they are, within the above mentioned constraint.
#     - The different question types will be grouped under their respective question type header names.
#     For example, if the heading is 'Fill in the Blanks', the questions below it must be related to fill in the blanks.
 
# Here's the instructions on how you should return the content of the worksheet:
#     - Return only the question type headings followed by the respective questions under it. **DO NOT** send anything extra apart from it.
#     - The output structure should be as demonstrated in the following example worksheet for physics:
#         Example:
#         Fill in the Blanks
#         1. Electric current is measured in ____
#         2. The SI unit of resistance is ____
 
#         Match the Following
#         3. Volt
#         4. Ampere
#         a) Unit of current
#         b) Unit of potential difference
 
#         Multiple Choice Questions
#         6. Which instrument measures electric current?
#         a) Voltmeter
#         b) Ammeter
#         c) Thermometer
#         d) Barometer
 
#         Short Answer Type Questions
#         8. Define electric current
#         9. What is Ohm's Law?
       
#         True or False
#         10. Ampere is the SI unit of resistance
#         11. Current is directly proportional to Voltage if Resistance is constant
# IMPORTANT:
#     USE `create_worksheet_tool` to create the worksheet word file and return the dictionary (as JSON) containing the following keys:
#     `response`: model response indicating the status of the worksheet creation
#     `agent_called`: name of the agent that created the worksheet
#     `filepath`: path to the created worksheet file
#     # **NEVER DISPLAY OR RETURN THE WORKSHEET FILE CONTENT TO THE USER**
# {school_guardrails}
# {local_language_policy}
 
# """
# ,
#     tools=[create_worksheet_tool],
# )

# import os
# import re
# from docx import Document
# from docx.shared import Pt
# from google.adk.agents import Agent
# from google.adk.tools.tool_context import ToolContext
# school_guardrails = """
# ## School Environment Guardrails
# - Do NOT use or allow any profanity, hate speech, or inappropriate language.
# - Do NOT generate or reference any content that is violent, sexual, discriminatory, or otherwise inappropriate for a school setting.
# - Always maintain a respectful, positive, and encouraging tone suitable for children and students.
# - Do NOT provide or suggest any unsafe, illegal, or harmful activities.
# - If a user asks for something inappropriate, politely refuse and remind them this is a school environment.
# """
# local_language_policy = """
# ## Local Language Policy
# - If the user asks a question in a local (non-English) language, reply in the same language but type your response in English script (transliteration).
# - Always match the user's language and tone, but keep the script in English.
# """
# def set_run(run, size=12, bold=False):
#     run.font.size = Pt(size)
#     run.bold = bold
#     run.font.name = "Calibri"
 
# # Helper to format subscripts and superscripts in a run (no LaTeX, no $...$)
# def add_formatted_text(paragraph, text):
#     import re
#     # Only handle superscript and subscript
#     pattern = re.compile(r'(\^\{.*?\}|\^[a-zA-Z0-9]|_\{.*?\}|_[a-zA-Z0-9])')
#     pos = 0
#     for m in pattern.finditer(text):
#         if m.start() > pos:
#             paragraph.add_run(text[pos:m.start()])
#         token = m.group()
#         if token.startswith('^'):
#             val = token[2:-1] if token.startswith('^{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.superscript = True
#         elif token.startswith('_'):
#             val = token[2:-1] if token.startswith('_{') else token[1]
#             run = paragraph.add_run(val)
#             run.font.subscript = True
#         pos = m.end()
#     if pos < len(text):
#         paragraph.add_run(text[pos:])
 
# def create_worksheet_tool(worksheet_text: str) -> dict:
#     """
#     This tool converts the contents of the worksheet to a word file
 
#     Args:
#         worksheet_text (str): contents of the worksheet
 
#     Returns:
#         dict:
#     """
#     if not worksheet_text:
#         return {"status": "Missing worksheet_text in context"}
 
#     lines = [line.strip() for line in worksheet_text.strip().splitlines() if line.strip()]
#     doc = Document()
#     section = None
#     match_left, match_right = [], []
 
#     i = 0
#     while i < len(lines):
#         line = lines[i]
#         # Section headers
#         if line in ["Fill in the Blanks", "Match the Following", "Multiple Choice Questions", "Short Answer Type Questions", "True or False","Solve These Problems"]:
#             section = line
#             p = doc.add_paragraph()
#             r = p.add_run(section)
#             set_run(r, size=14, bold=True)
#             i += 1
#             continue
 
#         if section == "Match the Following":
#             m_left = re.match(r'^(\d+)\.\s*(.+)$', line)
#             m_right = re.match(r'^([a-z])\)\s*(.+)$', line)
#             if m_left:
#                 match_left.append(m_left.group(2))
#             elif m_right:
#                 match_right.append(m_right.group(2))
#             if (i + 1 == len(lines) or lines[i + 1] in ["Multiple Choice Questions", "Short Answer Type Questions", "True or False"]):
#                 table = doc.add_table(rows=max(len(match_left), len(match_right)), cols=2)
#                 table.style = "Table Grid"
#                 for idx in range(len(match_left)):
#                     add_formatted_text(table.cell(idx, 0).paragraphs[0], f"{idx+1}. {match_left[idx]}")
#                     if idx < len(match_right):
#                         add_formatted_text(table.cell(idx, 1).paragraphs[0], f"{chr(97+idx)}) {match_right[idx]}")
#                 match_left, match_right = [], []
#             i += 1
#             continue
 
#         if section == "Multiple Choice Questions":
#             q_match = re.match(r'^(\d+)\.\s*(.+)$', line)
#             if q_match:
#                 question = q_match.group(2)
#                 options = []
#                 j = i + 1
#                 while j < len(lines) and re.match(r'^[a-d]\)', lines[j]):
#                     options.append(lines[j].strip())
#                     j += 1
#                 para = doc.add_paragraph()
#                 add_formatted_text(para, f"{q_match.group(1)}. {question}")
#                 opt_line = "    ".join(options)
#                 if len(opt_line) < 70:
#                     add_formatted_text(para, "\t" + opt_line)
#                 else:
#                     for opt in options:
#                         opt_p = doc.add_paragraph("\t")
#                         add_formatted_text(opt_p, opt)
#                 i = j
#                 continue
 
#         # Other questions
#         match_q = re.match(r'^(\d+)\.\s*(.+)$', line)
#         if match_q:
#             para = doc.add_paragraph()
#             add_formatted_text(para, f"{match_q.group(1)}. {match_q.group(2)}")
#         i += 1
 
#     base_path = r"C:\Users\birad\Downloads\Hack2skill_working_db\multi_agent\manager\sub_agents\worksheet"
#     target_path = os.path.join(base_path, "worksheets")
   
#     # Create the directory
#     os.makedirs(target_path, exist_ok=True)
#     filename = f"worksheets/worksheet_{hash(worksheet_text) % 100000}.docx"
#     doc.save(filename)
#     return {"status": "success", "worksheet_file": filename}
 
 
# root_agent = Agent(
#     name="worksheet_content_agent",
#     # https://ai.google.dev/gemini-api/docs/models
#     model="gemini-2.5-flash",
#     description="Creates content to be populated in the worksheet.",
#     instruction=f"""
# You are an expert at creating content of worksheets for students provided the class, syllabus and concept for creating the worksheet.
 
# The types of questions you must set will be as follows (these will apply for all subjects apart from Mathematics):
#     - Fill in the Blanks
#     - Match the Following
#     - Multiple Choice Questions
#     - Short Answer Type Questions
#     - True or False
# In case of Mathematics, do the following:
#     - type of questions:"solve these problems" and the content should be in the format of a worksheet.
#     - Create numerical problems of varying difficulties based on the syllabus and concept provided by the user
 
   
# Formatting instructions for worksheet content:
#     - For superscripts, use ^ (e.g., x^2; for multi-character superscripts, use curly brackets to enclose the entire superscript)
#     - For subscripts, use _ (e.g., H_2; for multi-character subscripts, use curly brackets to enclose the entire subscript)
#     - For equations, fractions, and roots, use plain text: e.g., 1/2, sqrt(x), x^2 + y^2 = 1
#     - For math symbols, use Unicode if possible (e.g., ≤, ≥, π, ×, ÷, ∪, ∩, ∈, ∉, ∞, etc.)
#     - For chemical symbols, use standard notation (e.g., H₂SO₄)
#     - Do NOT use LaTeX, $...$, or \frac, \leq, etc.
#     - Ensure formatting is consistent and clear for all questions
 
# Keep the following things in mind while creating worksheets:
#     - Make sure the class, syllabus and concept are mentioned before creating the worksheet. If any of these are missing, ask the user to mention them and do not proceed until they are properly passed.
#     - If not mentioned by the user, you will ask them how many questions per question type they expect and if they don't have a preference, default to 5 questions per question type
#     - You have to create questions only based on the specified class, syllabus and concept. No questions should appear beyond this constraint.
#     - Try to include questions of all difficulties (Easy, Medium, Hard) without mentioning which level they are, within the above mentioned constraint.
#     - The different question types will be grouped under their respective question type header names.
#     For example, if the heading is 'Fill in the Blanks', the questions below it must be related to fill in the blanks.
 
# Here's the instructions on how you should return the content of the worksheet:
#     - Return only the question type headings followed by the respective questions under it. **DO NOT** send anything extra apart from it.
#     - The output structure should be as demonstrated in the following example worksheet for physics:
#         Example:
#         Fill in the Blanks
#         1. Electric current is measured in ____
#         2. The SI unit of resistance is ____
 
#         Match the Following
#         3. Volt
#         4. Ampere
#         a) Unit of current
#         b) Unit of potential difference
 
#         Multiple Choice Questions
#         6. Which instrument measures electric current?
#         a) Voltmeter
#         b) Ammeter
#         c) Thermometer
#         d) Barometer
 
#         Short Answer Type Questions
#         8. Define electric current
#         9. What is Ohm's Law?
       
#         True or False
#         10. Ampere is the SI unit of resistance
#         11. Current is directly proportional to Voltage if Resistance is constant
# IMPORTANT:
#     USE `create_worksheet_tool` to create the worksheet word file and return the dictionary (as JSON) containing the following keys:
#     `response`: model response indicating the status of the worksheet creation
#     `agent_called`: name of the agent that created the worksheet
#     `filepath`: path to the created worksheet file
#     # **NEVER DISPLAY OR RETURN THE WORKSHEET FILE CONTENT TO THE USER**
# {school_guardrails}
# {local_language_policy}
 
# """
# ,
#     tools=[create_worksheet_tool],
# )